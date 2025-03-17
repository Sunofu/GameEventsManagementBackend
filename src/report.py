from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from fastapi import APIRouter, Request
from fastapi.responses import FileResponse, JSONResponse
from docx import Document
import os
from sqlalchemy import text
from backend.src.database import async_session
from datetime import datetime


report_router = APIRouter()

REPORT_FILE = "reports/subscription_report.docx"
LAST_REPORT_TIME = None
now = datetime.now()

@report_router.get("/")
async def get_report_page():
    return {
        "description": "Отчет по количеству подписок у каждой игры, самой популярной игре(-ах) и пользователях за текущий день",
        "last_report_time": LAST_REPORT_TIME if LAST_REPORT_TIME else "Отчет еще не был создан.",
    }

def set_heading_style(paragraph, level):
    """Установка стиля заголовков с отступами."""
    paragraph.style = f"Heading {level}"
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.space_after = Pt(12)

def add_table_header(table, headers):
    """Добавление заголовков для таблицы."""
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.bold = True

@report_router.post("/create")
async def create_report():
    global LAST_REPORT_TIME

    try:
        # Инициализация Word-документа
        doc = Document()
        doc.add_heading("Отчет по подпискам", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Дата создания отчета: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                          style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        async with async_session() as session:
            result_services = await session.execute(text("""
                SELECT g.title, COUNT(*) AS subscribes_count
                FROM game g
                JOIN subscriptions s ON g.id = s.gameid
                GROUP BY g.id
            """))

            doc.add_heading("Количество подписок у каждой игры", level=2)
            table1 = doc.add_table(rows=1, cols=2)
            table1.style = "Table Grid"
            add_table_header(table1, ["Игра", "Количество подписок"])

            for row in result_services:
                cells = table1.add_row().cells
                cells[0].text = row.title
                cells[1].text = str(row.subscribes_count)

            # 2. Самая популярная игра
            result_popular_game = await session.execute(text("""
                SELECT subquery.title, MAX(subquery.subscribes_count) AS subscribes_count
                FROM (
                    SELECT g.title, COUNT(*) AS subscribes_count
                    FROM game g
                    JOIN subscriptions s ON g.id = s.gameid
                    GROUP BY g.title
                ) AS subquery
                GROUP BY subquery.title
            """))
            popular_game = result_popular_game.fetchone()

            doc.add_heading("Самая популярная игра", level=2)
            if popular_game:
                para = doc.add_paragraph()
                para.add_run("Игра: ").bold = True
                para.add_run(f"{popular_game.title}").bold = True
                para.add_run(f" с количеством подписок: {popular_game.subscribes_count}")

            result_active_users = await session.execute(text("""
                SELECT username, email
                FROM users
                WHERE last_enter_date = CURRENT_DATE
            """))

            doc.add_heading("Пользователи за сегодня", level=2)
            table2 = doc.add_table(rows=1, cols=2)
            table2.style = "Table Grid"
            add_table_header(table2, ["Имя пользователя", "Почта пользователя"])

            for row in result_active_users:
                cells = table2.add_row().cells
                cells[0].text = row.username
                cells[1].text = str(row.email)

        os.makedirs("reports", exist_ok=True)
        doc.save(REPORT_FILE)

        LAST_REPORT_TIME = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        return JSONResponse({"status": "success", "message": "Отчет успешно создан"})

    except Exception as e:
        return JSONResponse({"status": "error", "message": str(e)})
    LAST_REPORT_TIME = now

@report_router.get("/download")
async def download_report():
    try:
        if not os.path.exists(REPORT_FILE):
            return JSONResponse({"status": "error", "message": "Отчет не найден"})
        return FileResponse(REPORT_FILE, filename="subscription_report.docx")
    except Exception as e:
        return JSONResponse({"status": "error", "message": str(e)})
